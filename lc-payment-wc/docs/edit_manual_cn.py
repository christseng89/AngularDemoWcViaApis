import copy
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table

PATH = "LC-Payment-WC-User-Manual-cn.docx"
d = docx.Document(PATH)
body = d.element.body
children = list(body)

# ---- 1. Fix stale counts / drop chargeBridge references ----

p109 = Paragraph(children[109], d)
assert '24 个业务案例' in p109.runs[0].text
p109.runs[0].text = (
    "本应用的第三个顶层标签页 —— Payment Component Simulator —— 与第 2–5 节是完全不同性质的功能："
    "它直接对接 microservices/payment-component 微服务，可从 23 个业务案例（每个都溯源自遗留系统"
    "的 Confirm 按钮函数，或对 RPFM 的部分整合案例而言，溯源自其分类步骤）中任选一个，编辑借/贷分录，"
    "并实时观察分类结果、余额校验、分录明细与 SWIFT 报文的重新计算。"
)

p112 = Paragraph(children[112], d)
assert p112.runs[0].text == '6.2 24 个业务案例'
p112.runs[0].text = '6.2 23 个业务案例'

t114 = Table(children[114], d)
cell_count = t114.rows[1].cells[1]
assert cell_count.paragraphs[0].runs[0].text == '16'
cell_count.paragraphs[0].runs[0].text = '15'
cell_meaning = t114.rows[1].cells[2]
assert '第 16 个' in cell_meaning.paragraphs[0].runs[0].text
cell_meaning.paragraphs[0].runs[0].text = (
    "已确认的 Payment Component 调用方，源代码中有可运作的分录组装流程 —— 模拟器会端到端驱动真实"
    "微服务（即时预览 + 真实 Confirm），这 15 个全部是源代码溯源验证过的。"
)

print("Step 1 (stale counts) done")
d.save(PATH)
